from docx import Document
import io
import base64

def extract_docx(file_bytes: bytes) -> dict:
    """
    Extract ALL content from DOCX including paragraphs, tables, headers, footers.
    Returns:
    {
      "blocks": [...],
      "original_file_base64": "base64_encoded_doc"
    }
    """
    doc = Document(io.BytesIO(file_bytes))
    blocks = []
    index = 0

    # Track paragraph positions for precise matching in exporter
    para_position = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            para_position += 1
            continue

        style_name = para.style.name.lower() if para.style else ''
        block_type = 'heading' if 'heading' in style_name else 'paragraph'

        blocks.append({
            "index": index,
            "type": block_type,
            "text": text,
            "element_index": para_position,  # Actual position in doc.paragraphs
            "element_type": "paragraph"
        })
        index += 1
        para_position += 1

    # Extract table cells - iterate ALL paragraphs inside each cell
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Collect text from ALL paragraphs in the cell
                cell_full_text = []
                for cp in cell.paragraphs:
                    t = cp.text.strip()
                    if t:
                        cell_full_text.append(t)
                
                full_text = "\n".join(cell_full_text)
                if not full_text:
                    continue

                blocks.append({
                    "index": index,
                    "type": "table_cell",
                    "text": full_text,
                    "element_index": index,
                    "element_type": "table_cell",
                    "table_idx": table_idx,
                    "row_idx": row_idx,
                    "col_idx": col_idx
                })
                index += 1

    # Extract headers & footers
    for s_idx, section in enumerate(doc.sections):
        for header_type in ['header', 'first_page_header', 'even_page_header']:
            header = getattr(section, header_type, None)
            if header:
                for p_idx, para in enumerate(header.paragraphs):
                    text = para.text.strip()
                    if text:
                        blocks.append({
                            "index": index,
                            "type": "header",
                            "text": text,
                            "element_index": f"header_{s_idx}_{header_type}_{p_idx}",
                            "element_type": "header"
                        })
                        index += 1

        for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
            footer = getattr(section, footer_type, None)
            if footer:
                for p_idx, para in enumerate(footer.paragraphs):
                    text = para.text.strip()
                    if text:
                        blocks.append({
                            "index": index,
                            "type": "footer",
                            "text": text,
                            "element_index": f"footer_{s_idx}_{footer_type}_{p_idx}",
                            "element_type": "footer"
                        })
                        index += 1

    return {
        "blocks": blocks,
        "original_file_base64": base64.b64encode(file_bytes).decode('utf-8')
    }
