from docx import Document
from docx.shared import Pt
import io
import base64

def export_docx(blocks: list[dict], original_file_base64: str) -> bytes:
    """
    Reconstruct DOCX by replacing text in the ORIGINAL file, preserving formatting.
    Uses exact structural keys (element_index for paragraphs, headers, footers) and 
    precise table cell coordinates for 100% deterministic replacement, regardless
    of spacing or split sentences.
    """
    original_bytes = base64.b64decode(original_file_base64)
    doc = Document(io.BytesIO(original_bytes))

    # 1. Group segments by structural container
    para_groups = {}   # { para_position_int : [segments_in_order] }
    table_groups = {}  # { "tableIdx_rowIdx_colIdx" : [segments_in_order] }
    header_groups = {} # { "header_S0_H_P0" : [segments_in_order] }
    footer_groups = {} # { "footer_S0_F_P0" : [segments_in_order] }
    
    for b in blocks:
        # Sort blocks into their respective structural groups by exact ID
        elem_type = b.get("element_type", "paragraph")
        
        if elem_type == "paragraph":
            idx = int(b.get("element_index", 0))
            para_groups.setdefault(idx, []).append(b)
            
        elif elem_type == "table_cell":
            t = b.get("table_idx", 0)
            r = b.get("row_idx", 0)
            c = b.get("col_idx", 0)
            key = f"{t}_{r}_{c}"
            table_groups.setdefault(key, []).append(b)
            
        elif elem_type == "header":
            idx = b.get("element_index", "")
            header_groups.setdefault(idx, []).append(b)
            
        elif elem_type == "footer":
            idx = b.get("element_index", "")
            footer_groups.setdefault(idx, []).append(b)

    import re
    # Helper: combine a group of segments by replacing them meticulously in the original paragraph text
    def _combine(group_blocks, original_text):
        if not original_text:
            return ""
            
        final_text = original_text
        sorted_blks = sorted(group_blocks, key=lambda x: x.get("index", 0) or x.get("original_block_index", 0))
        
        for bk in sorted_blks:
            orig = bk.get("text", "").strip()
            trans = bk.get("translated_text", "").strip()
            
            # If translation is empty or same, skip replacement
            if not orig or not trans or orig == trans:
                continue
                
            # Create a regex to match the original segment flexibly,
            # ignoring any spaces/newlines that might have been normalized during preprocessing
            escaped_orig = re.escape(orig)
            pattern = escaped_orig.replace(r'\ ', r'\s+')
            
            try:
                final_text = re.sub(pattern, lambda m: trans, final_text, count=1)
            except Exception as e:
                # If regex fails for some weird character combination, fallback to normal replace
                final_text = final_text.replace(orig, trans, 1)
                
        return final_text

    # 2. Apply Paragraph Replacements
    for i, para in enumerate(doc.paragraphs):
        if i in para_groups:
            new_text = _combine(para_groups[i], para.text)
            if new_text and new_text != para.text:
                _replace_para_text(para, new_text)

    # 3. Apply Table Replacements
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                key = f"{t_idx}_{r_idx}_{c_idx}"
                if key in table_groups:
                    orig_cell_text = "\\n".join([p.text for p in cell.paragraphs])
                    new_text = _combine(table_groups[key], orig_cell_text)
                    if not new_text or new_text == orig_cell_text:
                        continue
                        
                    # Apply to first paragraph's runs to preserve table format
                    if len(cell.paragraphs) > 0:
                        _replace_para_text(cell.paragraphs[0], new_text)
                        # Clear remaining paragraphs using the safe method to preserve internal drawings
                        for pi in range(1, len(cell.paragraphs)):
                            _replace_para_text(cell.paragraphs[pi], "")
                    else:
                        cell.text = new_text

    # 4. Apply Header Replacements
    for s_idx, section in enumerate(doc.sections):
        for header_type in ['header', 'first_page_header', 'even_page_header']:
            header = getattr(section, header_type, None)
            if header:
                for p_idx, para in enumerate(header.paragraphs):
                    key = f"header_{s_idx}_{header_type}_{p_idx}"
                    if key in header_groups:
                        new_text = _combine(header_groups[key], para.text)
                        if new_text and new_text != para.text:
                            _replace_para_text(para, new_text)

    # 5. Apply Footer Replacements
    for s_idx, section in enumerate(doc.sections):
        for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
            footer = getattr(section, footer_type, None)
            if footer:
                for p_idx, para in enumerate(footer.paragraphs):
                    key = f"footer_{s_idx}_{footer_type}_{p_idx}"
                    if key in footer_groups:
                        new_text = _combine(footer_groups[key], para.text)
                        if new_text and new_text != para.text:
                            _replace_para_text(para, new_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

from docx.oxml import OxmlElement

def _replace_para_text(para, new_text: str):
    """
    Replace all text in a paragraph by modifying only XML text nodes (w:t)
    while strictly preserving non-text sibling nodes (like w:drawing for images/logos).
    Also accurately maps \\n to w:br.
    """
    wt_elements = para._element.xpath('.//w:t')
    if not wt_elements:
        if new_text and new_text.strip():
            # If paragraph has NO text nodes but new text was sent, safely append
            para.add_run(new_text)
        return

    # Pick the text node that had the most characters originally to serve 
    # as the styling parent for new_text.
    best_wt = None
    max_len = -1
    for wt in wt_elements:
        t = wt.text or ""
        if len(t) > max_len:
            max_len = len(t)
            best_wt = wt
            
    if best_wt is None:
        best_wt = wt_elements[0]

    for wt in wt_elements:
        if wt == best_wt:
            _set_wt_text_with_newlines(wt, new_text)
        else:
            wt.text = ""

def _set_wt_text_with_newlines(wt, text: str):
    """
    Helper to inject new text safely with <w:br> support into a <w:t> element.
    """
    if not text:
        wt.text = ""
        return
        
    parts = text.split('\n')
    wt.text = parts[0]
    
    parent = wt.getparent()
    idx = parent.index(wt)
    
    for part in parts[1:]:
        br = OxmlElement('w:br')
        wt_new = OxmlElement('w:t')
        wt_new.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        wt_new.text = part
        
        idx += 1
        parent.insert(idx, br)
        idx += 1
        parent.insert(idx, wt_new)
